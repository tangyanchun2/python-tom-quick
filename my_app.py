from docx import Document
from docx.shared import Inches

import pyttsx3

def speak(text):
    pyttsx3.speak(text)


speak("hello, wbb")

#as dd
document = Document()
document.add_picture("girl5.jpg", width= Inches(2))
document.add_paragraph("this is chick 2 \n this is bad")
document.add_heading("it is goods")
p = document.add_paragraph("this is another chick 2")
p.add_run("this is another chick 2").bold =True

footer =document.sections[0].footer
p1 = footer.paragraphs[0]
p1.text ="haed heh"
document.save("my.docx")